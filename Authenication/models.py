# models.py
from django.db import models
from django.contrib.auth.models import AbstractBaseUser, BaseUserManager, PermissionsMixin
from django.core.validators import EmailValidator
from django.core.exceptions import ValidationError
from django.utils import timezone
import uuid
from django.conf import settings
from django.utils.text import slugify
import os
# Add these imports
import tempfile
import os
import pdfkit
from docx import Document
from io import BytesIO
import hashlib
from django.core.files.base import ContentFile
# Create your models here.
def sanitize_string(value):
    """
    Sanitize string input by removing potentially dangerous characters.
    """
    if not value:
        return value
    
    # Remove any null bytes
    value = value.replace('\x00', '')
    
    # Strip whitespace
    value = value.strip()
    
    # You can add more sanitization logic here as needed
    # For example, for email specific sanitization:
    if '@' in value:
        value = value.lower()
    
    return value

class CustomUserManager(BaseUserManager):
    def create_user(self, email, username, password=None, **extra_fields):
        """
        Create and return a regular user with email, username and password.
        """
        if not email:
            raise ValueError('Users must have an email address')
        if not username:
            raise ValueError('Users must have a username')
        
        email = self.normalize_email(email)
        user = self.model(email=email, username=username, **extra_fields)
        user.set_password(password)
        user.save(using=self._db)
        return user

    def create_superuser(self, email, username, password=None, **extra_fields):
        """
        Create and return a superuser with email, username and password.
        """
        extra_fields.setdefault('is_staff', True)
        extra_fields.setdefault('is_superuser', True)
        extra_fields.setdefault('is_active', True)
        extra_fields.setdefault('user_type', 'admin')  # Add this line
        
        if extra_fields.get('is_staff') is not True:
            raise ValueError('Superuser must have is_staff=True.')
        if extra_fields.get('is_superuser') is not True:
            raise ValueError('Superuser must have is_superuser=True.')
        
        return self.create_user(email, username, password, **extra_fields)

class User(AbstractBaseUser, PermissionsMixin):
    """
    Custom User model that extends Django's AbstractBaseUser and PermissionsMixin.
    """
    id = models.UUIDField(
        primary_key=True,
        default=uuid.uuid4,
        editable=False,
        verbose_name='ID'
    )
    username = models.CharField(max_length=150, unique=True)
    email = models.EmailField(
        unique=True,
        validators=[EmailValidator()],
        verbose_name='email address'
    )
    
    # Additional fields
    first_name = models.CharField(max_length=30, blank=True)
    last_name = models.CharField(max_length=30, blank=True)
    date_joined = models.DateTimeField(default=timezone.now)
    
    # Required fields for Django admin
    is_staff = models.BooleanField(default=False)
    is_active = models.BooleanField(default=True)
    is_superuser = models.BooleanField(default=False)
    bio = models.TextField(
        max_length=500,
        blank=True,
        null=True,
        help_text='User biography or description'
    )
    
    # User types (add this field - was missing!)
    USER_TYPE_CHOICES = (
        ('admin', 'Administrator'),
        ('user', 'Regular User'),
        ('moderator', 'Moderator'),
    )
    user_type = models.CharField(
        max_length=20,
        choices=USER_TYPE_CHOICES,
        default='user'
    )
    
    # Track timestamps
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    
    # Define the manager
    objects = CustomUserManager()
    
    # Specify which field should be used as the unique identifier
    USERNAME_FIELD = 'email'
    REQUIRED_FIELDS = ['username']
    
    class Meta:
        verbose_name = 'User'
        verbose_name_plural = 'Users'
        ordering = ['-date_joined']
    
    def __str__(self):
        return self.username
    
    def get_full_name(self):
        """
        Return the first_name plus the last_name, with a space in between.
        """
        full_name = f'{self.first_name} {self.last_name}'
        return full_name.strip()
    
    def get_short_name(self):
        """
        Return the short name for the user.
        """
        return self.first_name
    
    @property
    def hex_id(self):
        """Get 16-character hex representation of UUID."""
        return str(self.id).replace('-', '')[:16].lower()
    
    @property
    def is_admin(self):
        """Check if user is an admin."""
        return self.user_type == 'admin' or self.is_superuser
    
    @property
    def is_moderator(self):

        """Check if user is a moderator."""
        return self.user_type == 'moderator'
    

def validate_pdf_or_docx(value):
    """
    Validate that uploaded file is either PDF or DOCX.
    """
    ext = os.path.splitext(value.name)[1].lower()
    valid_extensions = ['.pdf', '.docx', '.doc']
    if not ext in valid_extensions:
        raise ValidationError('Unsupported file format. Please upload PDF or DOCX files only.')

def course_upload_path(instance, filename):
    """
    Generate upload path for course files.
    Format: courses/{course_slug}/{filename}
    """
    ext = filename.split('.')[-1]
    filename = f"course_file.{ext}"
    return f'courses/{instance.slug}/{filename}'

def lesson_upload_path(instance, filename):
    """
    Generate upload path for lesson files.
    Format: courses/{course_slug}/lessons/{lesson_slug}/{filename}
    """
    ext = filename.split('.')[-1]
    filename = f"lesson_file.{ext}"
    return f'courses/{instance.course.slug}/lessons/{instance.slug}/{filename}'

class Course(models.Model):
    """
    Model representing a course created by a user.
    """
    id = models.UUIDField(
        primary_key=True,
        default=uuid.uuid4,
        editable=False,
        verbose_name='ID'
    )
    
    # Basic course information
    title = models.CharField(max_length=200)
    slug = models.SlugField(max_length=250, unique=True, blank=True)
    description = models.TextField(blank=True)
    
    # Course creator (owner)
    creator = models.ForeignKey(
        settings.AUTH_USER_MODEL,
        on_delete=models.CASCADE,
        related_name='created_courses',
        verbose_name='Course Creator'
    )
    
    # Course file (PDF or DOCX)
    course_file = models.FileField(
        upload_to=course_upload_path,
        validators=[validate_pdf_or_docx],
        help_text='Upload PDF or DOCX file (max 50MB)'
    )
    
    # File metadata
    file_size = models.PositiveIntegerField(default=0, editable=False)  # in bytes
    file_type = models.CharField(max_length=10, editable=False)  # pdf or docx
    total_pages = models.PositiveIntegerField(default=0, editable=False)
    
    # Course settings
    is_public = models.BooleanField(default=False)
    allow_comments = models.BooleanField(default=True)
    enable_progress_tracking = models.BooleanField(default=True)
    
    # Timestamps
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    published_at = models.DateTimeField(null=True, blank=True)
    
    # Course statistics (can be updated via signals or methods)
    total_lessons = models.PositiveIntegerField(default=0, editable=False)
    total_enrollments = models.PositiveIntegerField(default=0, editable=False)
    average_rating = models.FloatField(default=0.0, editable=False)
    
    class Meta:
        verbose_name = 'Course'
        verbose_name_plural = 'Courses'
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['slug']),
            models.Index(fields=['creator', 'created_at']),
            models.Index(fields=['is_public', 'published_at']),
        ]
    
    def __str__(self):
        return self.title
    
    def save(self, *args, **kwargs):
        """
        Override save method to:
        1. Generate slug from title
        2. Extract file metadata
        3. Convert DOCX to PDF if needed
        """
        if not self.slug:
            self.slug = slugify(self.title)
            # Ensure slug is unique
            original_slug = self.slug
            counter = 1
            while Course.objects.filter(slug=self.slug).exists():
                self.slug = f'{original_slug}-{counter}'
                counter += 1
        
        is_new = self.pk is None
        original_file = None
        
        # If file is uploaded and it's a DOCX file
        if self.course_file and hasattr(self.course_file, 'size'):
            # Store original file for conversion
            if is_new:
                original_file = self.course_file
                
            # Determine file type
            ext = os.path.splitext(self.course_file.name)[1].lower()
            
            if ext == '.pdf':
                self.file_type = 'pdf'
                self.total_pages = self.get_pdf_page_count()
            elif ext in ['.docx', '.doc']:
                self.file_type = 'pdf'  # We'll convert to PDF
                # We'll handle conversion after initial save
            
            self.file_size = self.course_file.size
        
        # First save to get an ID
        super().save(*args, **kwargs)
        
        # Convert DOCX to PDF if needed (only for new uploads)
        if is_new and original_file and original_file.name.lower().endswith(('.docx', '.doc')):
            try:
                self.convert_docx_to_pdf(original_file)
                # Update file metadata after conversion
                self.total_pages = self.get_pdf_page_count()
                self.save(update_fields=['course_file', 'file_size', 'file_type', 'total_pages'])
            except Exception as e:
                print(f"Error converting DOCX to PDF: {e}")
                # Keep the original DOCX if conversion fails
                self.file_type = 'docx'
                self.total_pages = self.get_docx_page_count()
                self.save(update_fields=['file_type', 'total_pages'])
            super().save(*args, **kwargs)
    
    def get_pdf_page_count(self):
        """
        Count pages in PDF file.
        Note: You'll need to install PyPDF2: pip install PyPDF2
        """
        try:
            import PyPDF2
            from io import BytesIO
            
            # Read the PDF file
            self.course_file.seek(0)
            pdf_reader = PyPDF2.PdfReader(self.course_file)
            return len(pdf_reader.pages)
        except Exception as e:
            print(f"Error counting PDF pages: {e}")
            return 0
    
    def convert_docx_to_pdf(self, docx_file):
        """
        Convert DOCX file to PDF and update the course_file field.
        """
        try:
            # Read DOCX content
            docx_file.seek(0)
            doc = Document(docx_file)
            
            # Create HTML content from DOCX
            html_content = self.docx_to_html(doc)
            
            # Create PDF from HTML
            pdf_options = {
                'page-size': 'A4',
                'margin-top': '0.75in',
                'margin-right': '0.75in',
                'margin-bottom': '0.75in',
                'margin-left': '0.75in',
                'encoding': "UTF-8",
                'no-outline': None,
                'enable-local-file-access': None,
                'quiet': ''
            }
            
            # Configure pdfkit path (adjust for your system)
            config = pdfkit.configuration(wkhtmltopdf='C:/Program Files/wkhtmltopdf/bin/wkhtmltopdf.exe')  # Windows
            # config = pdfkit.configuration(wkhtmltopdf='/usr/local/bin/wkhtmltopdf')  # Mac/Linux
            
            # Generate PDF
            pdf_bytes = pdfkit.from_string(html_content, False, options=pdf_options, configuration=config)
            
            # Create new filename with PDF extension
            original_name = os.path.splitext(docx_file.name)[0]
            new_filename = f"{original_name}.pdf"
            
            # Save PDF to course_file field
            self.course_file.save(
                new_filename,
                ContentFile(pdf_bytes),
                save=False
            )
            
            # Update file size
            self.file_size = len(pdf_bytes)
            
        except Exception as e:
            print(f"Error in DOCX to PDF conversion: {e}")
            raise

    def docx_to_html(self, doc):
        """
        Convert DOCX document to HTML for PDF conversion.
        """
        html_parts = []
        
        # Start HTML document
        html_parts.append("""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body { font-family: Arial, sans-serif; line-height: 1.6; margin: 0; padding: 20px; }
                h1 { font-size: 24px; margin-top: 20px; margin-bottom: 10px; }
                h2 { font-size: 20px; margin-top: 18px; margin-bottom: 8px; }
                h3 { font-size: 16px; margin-top: 16px; margin-bottom: 6px; }
                p { margin-bottom: 10px; }
                ul, ol { margin-left: 20px; margin-bottom: 10px; }
                table { border-collapse: collapse; width: 100%; margin-bottom: 10px; }
                th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
                th { background-color: #f2f2f2; }
                .page-break { page-break-after: always; }
            </style>
        </head>
        <body>
        """)
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Skip empty paragraphs
                style = paragraph.style.name
                
                if style.startswith('Heading'):
                    level = int(style[-1]) if style[-1].isdigit() else 1
                    html_parts.append(f"<h{level}>{paragraph.text}</h{level}>")
                else:
                    html_parts.append(f"<p>{paragraph.text}</p>")
        
        # Process tables
        for table in doc.tables:
            html_parts.append("<table>")
            for row in table.rows:
                html_parts.append("<tr>")
                for cell in row.cells:
                    html_parts.append(f"<td>{cell.text}</td>")
                html_parts.append("</tr>")
            html_parts.append("</table>")
        
        # End HTML document
        html_parts.append("""
        </body>
        </html>
        """)
        
        return '\n'.join(html_parts)
    
    def get_docx_page_count(self):
        """
        Estimate pages in DOCX file.
        Note: You'll need to install python-docx: pip install python-docx
        """
        try:
            from docx import Document
            from io import BytesIO
            
            self.course_file.seek(0)
            doc = Document(self.course_file)
            
            # Estimate pages based on paragraphs (rough estimate)
            # You might want to implement more accurate counting
            total_paragraphs = len(doc.paragraphs)
            estimated_pages = max(1, total_paragraphs // 10)  # Rough estimate
            return estimated_pages
        except Exception as e:
            print(f"Error estimating DOCX pages: {e}")
            return 0
    
    def get_absolute_url(self):
        from django.urls import reverse
        return reverse('course-detail', kwargs={'slug': self.slug})
    
    @property
    def total_duration(self):
        """Calculate total estimated duration in minutes."""
        # Estimate 5 minutes per page for reading
        return self.total_pages * 5
    
    @property
    def is_published(self):
        """Check if course is published."""
        return self.is_public and self.published_at is not None
    
    def enroll_user(self, user):
        """Enroll a user in this course."""
        enrollment, created = CourseEnrollment.objects.get_or_create(
            user=user,
            course=self,
            defaults={'enrolled_at': timezone.now()}
        )
        if created:
            self.total_enrollments += 1
            self.save()
        return enrollment


class Lesson(models.Model):
    """
    Model representing individual lessons/sections within a course.
    Each lesson corresponds to a specific page or section of the document.
    """
    id = models.UUIDField(
        primary_key=True,
        default=uuid.uuid4,
        editable=False,
        verbose_name='ID'
    )
    
    course = models.ForeignKey(
        Course,
        on_delete=models.CASCADE,
        related_name='lessons'
    )
    
    # Lesson information
    title = models.CharField(max_length=200)
    slug = models.SlugField(max_length=250, blank=True)
    description = models.TextField(blank=True)
    
    # Page/section information
    page_number = models.PositiveIntegerField(
        help_text='Page number in the document'
    )
    page_range = models.CharField(
        max_length=50,
        blank=True,
        help_text='Page range (e.g., "1-5" or "Chapter 1")'
    )
    
    # Duration estimate
    estimated_duration = models.PositiveIntegerField(
        default=0,
        help_text='Estimated duration in minutes'
    )
    
    # Lesson content (optional - can have supplementary files)
    lesson_file = models.FileField(
        upload_to=lesson_upload_path,
        validators=[validate_pdf_or_docx],
        blank=True,
        null=True,
        help_text='Optional supplementary file for this lesson'
    )
    
    # Ordering
    order = models.PositiveIntegerField(
        default=0,
        help_text='Order of lesson in the course'
    )
    
    # Timestamps
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    
    class Meta:
        verbose_name = 'Lesson'
        verbose_name_plural = 'Lessons'
        ordering = ['order', 'page_number']
        unique_together = ['course', 'page_number']
        indexes = [
            models.Index(fields=['course', 'order']),
            models.Index(fields=['course', 'page_number']),
        ]
    
    def __str__(self):
        return f"{self.course.title} - {self.title} (Page {self.page_number})"
    
    def save(self, *args, **kwargs):
        """Generate slug from title."""
        if not self.slug:
            self.slug = slugify(self.title)
        super().save(*args, **kwargs)
    
    @property
    def next_lesson(self):
        """Get the next lesson in sequence."""
        try:
            return Lesson.objects.filter(
                course=self.course,
                order__gt=self.order
            ).order_by('order').first()
        except Lesson.DoesNotExist:
            return None
    
    @property
    def previous_lesson(self):
        """Get the previous lesson in sequence."""
        try:
            return Lesson.objects.filter(
                course=self.course,
                order__lt=self.order
            ).order_by('-order').first()
        except Lesson.DoesNotExist:
            return None


class CourseEnrollment(models.Model):
    """
    Model representing user enrollment in a course.
    """
    id = models.UUIDField(
        primary_key=True,
        default=uuid.uuid4,
        editable=False,
        verbose_name='ID'
    )
    
    user = models.ForeignKey(
        settings.AUTH_USER_MODEL,
        on_delete=models.CASCADE,
        related_name='enrollments'
    )
    
    course = models.ForeignKey(
        Course,
        on_delete=models.CASCADE,
        related_name='enrollments'
    )
    
    enrolled_at = models.DateTimeField(auto_now_add=True)
    completed_at = models.DateTimeField(null=True, blank=True)
    
    # Current progress
    current_page = models.PositiveIntegerField(default=1)
    current_lesson = models.ForeignKey(
        Lesson,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='active_enrollments'
    )
    
    # Progress statistics
    pages_completed = models.PositiveIntegerField(default=0)
    total_pages_viewed = models.PositiveIntegerField(default=0)
    last_accessed = models.DateTimeField(auto_now=True)
    
    # User preferences
    bookmarked_pages = models.JSONField(
        default=list,
        help_text='List of bookmarked page numbers'
    )
    notes = models.JSONField(
        default=dict,
        help_text='Notes for specific pages: {"page_number": "note_text"}'
    )
    
    class Meta:
        verbose_name = 'Course Enrollment'
        verbose_name_plural = 'Course Enrollments'
        unique_together = ['user', 'course']
        ordering = ['-last_accessed']
        indexes = [
            models.Index(fields=['user', 'course']),
            models.Index(fields=['user', 'last_accessed']),
            models.Index(fields=['course', 'completed_at']),
        ]
    
    def __str__(self):
        return f"{self.user.username} - {self.course.title}"
    
    @property
    def is_completed(self):
        """Check if course is completed."""
        return self.completed_at is not None
    
    @property
    def progress_percentage(self):
        """Calculate progress percentage."""
        if self.course.total_pages > 0:
            return min(100, (self.pages_completed / self.course.total_pages) * 100)
        return 0
    
    @property
    def time_spent(self):
        """Calculate total time spent on course (in minutes)."""
        # This could be calculated from UserProgressLog entries
        from django.db.models import Sum
        total_time = UserProgressLog.objects.filter(
            enrollment=self
        ).aggregate(total=Sum('time_spent'))['total']
        return total_time or 0
    
    def update_progress(self, page_number, time_spent=0):
        """
        Update user's progress in the course.
        
        Args:
            page_number: The page the user is currently viewing
            time_spent: Time spent on the page (in seconds)
        """
        # Update current page
        old_page = self.current_page
        self.current_page = page_number
        
        # Update pages completed (if user viewed a new page)
        if page_number > self.pages_completed:
            self.pages_completed = page_number
        
        # Update current lesson
        try:
            self.current_lesson = Lesson.objects.get(
                course=self.course,
                page_number=page_number
            )
        except Lesson.DoesNotExist:
            # If no lesson exists for this page, find the closest one
            lesson = Lesson.objects.filter(
                course=self.course,
                page_number__lte=page_number
            ).order_by('-page_number').first()
            self.current_lesson = lesson
        
        # Mark as completed if user reached the last page
        if page_number >= self.course.total_pages and not self.completed_at:
            self.completed_at = timezone.now()
        
        self.save()
        
        # Log this progress update
        UserProgressLog.objects.create(
            enrollment=self,
            page_number=page_number,
            time_spent=time_spent,
            action='page_view'
        )
        
        # Update total pages viewed
        self.total_pages_viewed += 1
        self.save(update_fields=['total_pages_viewed'])
        
        return self


class UserProgressLog(models.Model):
    """
    Detailed log of user progress for analytics.
    """
    ACTION_CHOICES = [
        ('page_view', 'Page View'),
        ('quiz_complete', 'Quiz Completed'),
        ('bookmark', 'Bookmark Added'),
        ('note_added', 'Note Added'),
        ('course_complete', 'Course Completed'),
    ]
    
    id = models.UUIDField(
        primary_key=True,
        default=uuid.uuid4,
        editable=False,
        verbose_name='ID'
    )
    
    enrollment = models.ForeignKey(
        CourseEnrollment,
        on_delete=models.CASCADE,
        related_name='progress_logs'
    )
    
    # Progress details
    page_number = models.PositiveIntegerField()
    action = models.CharField(max_length=50, choices=ACTION_CHOICES)
    time_spent = models.PositiveIntegerField(
        default=0,
        help_text='Time spent in seconds'
    )
    
    # Additional data
    metadata = models.JSONField(
        default=dict,
        help_text='Additional data about the action'
    )
    
    created_at = models.DateTimeField(auto_now_add=True)
    
    class Meta:
        verbose_name = 'User Progress Log'
        verbose_name_plural = 'User Progress Logs'
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['enrollment', 'created_at']),
            models.Index(fields=['action', 'created_at']),
        ]
    
    def __str__(self):
        return f"{self.enrollment.user.username} - {self.action} at {self.created_at}"


class CourseRating(models.Model):
    """
    Model for user ratings and reviews of courses.
    """
    id = models.UUIDField(
        primary_key=True,
        default=uuid.uuid4,
        editable=False,
        verbose_name='ID'
    )
    
    user = models.ForeignKey(
        settings.AUTH_USER_MODEL,
        on_delete=models.CASCADE,
        related_name='course_ratings'
    )
    
    course = models.ForeignKey(
        Course,
        on_delete=models.CASCADE,
        related_name='ratings'
    )
    
    rating = models.PositiveSmallIntegerField(
        choices=[(i, str(i)) for i in range(1, 6)],
        help_text='Rating from 1 to 5 stars'
    )
    
    review = models.TextField(blank=True)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    
    class Meta:
        verbose_name = 'Course Rating'
        verbose_name_plural = 'Course Ratings'
        unique_together = ['user', 'course']
        ordering = ['-created_at']
    
    def __str__(self):
        return f"{self.user.username} - {self.course.title} ({self.rating} stars)"